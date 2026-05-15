from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers de estilo ─────────────────────────────────────────────────────────
def set_color(run, hex_color):
    r, g, b = int(hex_color[0:2],16), int(hex_color[2:4],16), int(hex_color[4:6],16)
    run.font.color.rgb = RGBColor(r, g, b)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ('top','left','bottom','right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'BFBFBF')
        tcBorders = OxmlElement('w:tcBorders')
        tcBorders.append(border)
        tcPr.append(tcBorders)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    set_color(run, "1F4E79")

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    set_color(run, "555555")

def add_h1(text):
    p = doc.add_heading(level=1)
    p.clear()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    set_color(run, "1F4E79")

def add_h2(text):
    p = doc.add_heading(level=2)
    p.clear()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    set_color(run, "2E75B6")

def add_p(text="", space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        p.add_run(text).font.size = Pt(11)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text).font.size = Pt(11)

def add_label(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    set_color(run, "1F4E79")

def add_json(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(4)
    # fondo gris claro via shading en el párrafo
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    set_color(run, "C00000")

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("↳  " + text)
    run.italic = True
    run.font.size = Pt(10)
    set_color(run, "555555")

def add_meta(topic, emisor, receptor, cuando):
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    col_w = [Cm(3), Cm(14)]
    for i, row in enumerate(tbl.rows):
        for j, cell in enumerate(row.cells):
            cell.width = col_w[j]
            shade_cell(cell, "EBF3FB" if j == 0 else "FFFFFF")
    data = [("Topic", topic), ("Emisor", emisor), ("Receptor", receptor), ("Cuándo", cuando)]
    for i, (k, v) in enumerate(data):
        kc = tbl.rows[i].cells[0]
        vc = tbl.rows[i].cells[1]
        kr = kc.paragraphs[0].add_run(k)
        kr.bold = True; kr.font.size = Pt(10)
        vr = vc.paragraphs[0].add_run(v)
        vr.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_summary_table(headers, rows, col_widths):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    hdr_cells = tbl.rows[0].cells
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        hdr_cells[i].width = Cm(w)
        shade_cell(hdr_cells[i], "BDD7EE")
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = tbl.add_row().cells
        for i, (val, w) in enumerate(zip(row, col_widths)):
            cells[i].width = Cm(w)
            p = cells[i].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if i == 2:
                run.font.name = "Courier New"
                set_color(run, "C00000")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def msg_block(topic, emisor, receptor, cuando, plantilla, ejemplo, accion):
    add_meta(topic, emisor, receptor, cuando)
    add_label("Plantilla:")
    add_json(plantilla)
    add_label("Ejemplo:")
    add_json(ejemplo)
    add_note(accion)

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ═══════════════════════════════════════════════════════════════════════════════

add_title("Protocolo de integración MQTT")
add_subtitle("Perspectiva RoboDK  ·  Proyecto GIIROB PR2-A1")

add_p("Este documento describe todos los mensajes MQTT que el entorno RoboDK (simulación Python) "
      "debe recibir y publicar para integrarse con el ESP32-S3 (controlador central), "
      "el robot Delta, el Cobot y el sistema de emergencia.")

# ─── 1. RESPONSABILIDADES ─────────────────────────────────────────────────────
add_h1("1. Responsabilidades de RoboDK")
add_bullet("Cámara virtual  —  detecta tapas spawneadas y publica coordenadas + color.")
add_bullet("Robot Delta     —  recibe órdenes de pick y mueve tapas a las tolvas.")
add_bullet("Cobot           —  recibe órdenes de paletizado y publica confirmación.")
add_bullet("Emergencia      —  escucha el estado y detiene / reanuda operaciones.")

# ─── 2. MENSAJES RECIBIDOS ────────────────────────────────────────────────────
add_h1("2. Mensajes que RoboDK RECIBE")

add_h2("2.1  Spawn de tapa — generación en escena")
msg_block(
    topic    = "giirob/pr2-A1/devices/robodk/action",
    emisor   = "ESP32-S3",
    receptor = "Bridge Python (RoboDK)",
    cuando   = "Cada vez que el ESP32 necesita generar una tapa nueva en la simulación.",
    plantilla= '{"cmd":"spawn", "id_cap":"<id>", "color":"<color>"}',
    ejemplo  = '{"cmd":"spawn", "id_cap":"C0042", "color":"blue"}',
    accion   = (
        "Copiar/pegar el objeto de tapa en la escena RoboDK con el color indicado. "
        "Una vez visible, publicar la detección en camera/data con el mismo id_cap."
    )
)

add_h2("2.2  Orden de pick — Robot Delta")
msg_block(
    topic    = "giirob/pr2-A1/devices/delta/action",
    emisor   = "ESP32-S3",
    receptor = "Bridge Python (Delta)",
    cuando   = "Cuando el ESP32 valida una tapa detectada y decide clasificarla.",
    plantilla= '{"cmd":"pick", "x":<x>, "y":<y>, "color":"<color>", "tolva":"TOLVA_<N>", "id_cap":"<id>", "reason":"<motivo>"}',
    ejemplo  = '{"cmd":"pick", "x":1.2, "y":3.4, "color":"red", "tolva":"TOLVA_1", "id_cap":"C0001", "reason":"Auto: aceptando tapa color red"}',
    accion   = (
        "Mover el robot Delta a las coordenadas (x, y), recoger la tapa y depositarla "
        "en la tolva indicada. Las tolvas siempre en mayúsculas: TOLVA_1 … TOLVA_6."
    )
)

add_h2("2.3  Orden de paletizado — Cobot")
msg_block(
    topic    = "giirob/pr2-A1/devices/cobot/action",
    emisor   = "ESP32-S3",
    receptor = "Cobot",
    cuando   = "Cuando el AMR llega a cobot_pick y el ESP32 autoriza el paletizado.",
    plantilla= '{"cmd":"start", "id_pallet":"<id>", "color":"<color>", "boxes_stacked":<n>}',
    ejemplo  = '{"cmd":"start", "id_pallet":"P0001", "color":"red", "boxes_stacked":3}',
    accion   = (
        "Recoger la caja del AMR y depositarla en el pallet. "
        "boxes_stacked indica cuántas cajas hay en ese pallet ANTES de esta operación. "
        "Al finalizar, publicar confirmación en cobot/status."
    )
)

add_h2("2.4  Emergencia activa")
msg_block(
    topic    = "giirob/pr2-A1/system/emergency/status",
    emisor   = "ESP32-S3",
    receptor = "Todos (Delta, Cobot, Bridge Python)",
    cuando   = "Al presionar botón GPIO38 o recibir estop por MQTT.",
    plantilla= '{"status":"emergency_active", "device":"<dispositivo>", "sensor":"<origen>"}',
    ejemplo  = '{"status":"emergency_active", "device":"ESP32-S3", "sensor":"emergency_button"}',
    accion   = (
        "Detener inmediatamente todos los movimientos del Delta y el Cobot. "
        "Vaciar la cola de picks pendientes. No iniciar ninguna operación nueva."
    )
)

add_h2("2.5  Emergencia inactiva — reanudación")
msg_block(
    topic    = "giirob/pr2-A1/system/emergency/status",
    emisor   = "ESP32-S3",
    receptor = "Todos (Delta, Cobot, Bridge Python)",
    cuando   = "Al presionar botón GPIO39 o recibir resume por MQTT.",
    plantilla= '{"status":"emergency_inactive", "device":"<dispositivo>", "sensor":"<origen>"}',
    ejemplo  = '{"status":"emergency_inactive", "device":"ESP32-S3", "sensor":"resume_button"}',
    accion   = "Reanudar las operaciones. El sistema vuelve a aceptar órdenes de spawn y pick."
)

# ─── 3. MENSAJES PUBLICADOS ───────────────────────────────────────────────────
add_h1("3. Mensajes que RoboDK PUBLICA")

add_h2("3.1  Detección de tapa — Cámara virtual")
msg_block(
    topic    = "giirob/pr2-A1/devices/camera/data",
    emisor   = "Bridge Python (cámara virtual)",
    receptor = "ESP32-S3",
    cuando   = "Inmediatamente después de crear la tapa en escena (tras spawn).",
    plantilla= '{"x":<x>, "y":<y>, "color":"<color>", "precision":<valor>, "id_cap":"<id>"}',
    ejemplo  = '{"x":1.2, "y":3.4, "color":"red", "precision":0.97, "id_cap":"C0001"}',
    accion   = (
        "El ESP32 procesa la detección solo si precision > 0.95. "
        "El id_cap debe ser el mismo recibido en el spawn correspondiente."
    )
)

add_h2("3.2  Confirmación de paletizado — Cobot")
msg_block(
    topic    = "giirob/pr2-A1/devices/cobot/status",
    emisor   = "Cobot",
    receptor = "ESP32-S3",
    cuando   = "Al finalizar el movimiento de depositar la caja en el pallet.",
    plantilla= '{"status":"completed", "id_pallet":"<id>"}',
    ejemplo  = '{"status":"completed", "id_pallet":"P0001"}',
    accion   = (
        "El ESP32 actualiza el contador del pallet. Si alcanza 12 cajas, cierra el pallet, "
        "solicita operario y notifica al SCADA. El id_pallet debe ser el mismo recibido en start."
    )
)

add_h2("3.3  Parada de emergencia desde Delta o Cobot")
msg_block(
    topic    = "giirob/pr2-A1/system/emergency/action",
    emisor   = "Delta / Cobot",
    receptor = "ESP32-S3",
    cuando   = "Si el Delta o el Cobot detectan un fallo grave durante la operación.",
    plantilla= '{"cmd":"estop", "source":"<dispositivo>", "reason":"<motivo>"}',
    ejemplo  = '{"cmd":"estop", "source":"DELTA", "reason":"collision"}',
    accion   = (
        "El ESP32 activará la parada de emergencia y publicará emergency_active a todos los dispositivos."
    )
)

# ─── 4. RESUMEN DE TOPICS ─────────────────────────────────────────────────────
add_h1("4. Resumen de topics")
add_summary_table(
    headers    = ["Topic", "RoboDK", "Mensaje clave"],
    rows       = [
        ["giirob/pr2-A1/devices/robodk/action",   "RECIBE",  '{"cmd":"spawn","id_cap":"C0042","color":"blue"}'],
        ["giirob/pr2-A1/devices/delta/action",    "RECIBE",  '{"cmd":"pick","x":1.2,"y":3.4,"color":"red","tolva":"TOLVA_1","id_cap":"C0001"}'],
        ["giirob/pr2-A1/devices/cobot/action",    "RECIBE",  '{"cmd":"start","id_pallet":"P0001","color":"red","boxes_stacked":3}'],
        ["giirob/pr2-A1/system/emergency/status", "RECIBE",  '{"status":"emergency_active / emergency_inactive","device":"ESP32-S3","sensor":"..."}'],
        ["giirob/pr2-A1/devices/camera/data",     "PUBLICA", '{"x":1.2,"y":3.4,"color":"red","precision":0.97,"id_cap":"C0001"}'],
        ["giirob/pr2-A1/devices/cobot/status",    "PUBLICA", '{"status":"completed","id_pallet":"P0001"}'],
        ["giirob/pr2-A1/system/emergency/action", "PUBLICA", '{"cmd":"estop","source":"DELTA","reason":"collision"}'],
    ],
    col_widths = [8.5, 2.5, 8.0]
)

# ─── 5. NOTAS ─────────────────────────────────────────────────────────────────
add_h1("5. Notas de integración")
add_bullet("El id_cap lo genera el ESP32 y se envía en spawn. El bridge Python debe incluir ese mismo id_cap en camera/data.")
add_bullet("Las tolvas siempre en mayúsculas: TOLVA_1 … TOLVA_6.")
add_bullet("El id_pallet crece indefinidamente: P0001, P0002, … sin límite de pallets.")
add_bullet("boxes_stacked es el conteo de cajas en el pallet ANTES de la operación actual (0 si es la primera).")
add_bullet("El ESP32 descarta detecciones con precision ≤ 0.95.")
add_bullet("En emergencia el bridge debe vaciar la cola de picks y no iniciar nuevos movimientos hasta recibir emergency_inactive.")

doc.save("c:/p/c/RoboDK_Protocolo_MQTT.docx")
print("Generado: RoboDK_Protocolo_MQTT.docx")
